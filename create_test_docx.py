from docx import Document

# Create a new DOCX document
doc = Document()

# Add title
doc.add_heading('RENTAL AGREEMENT', 0)

# Add content
doc.add_paragraph('This Rental Agreement is made between John Doe (Landlord) and Jane Smith (Tenant).')

doc.add_heading('PROPERTY DETAILS:', level=1)
doc.add_paragraph('Address: 123 Main Street, Mumbai, Maharashtra')

doc.add_heading('RENTAL TERMS:', level=1)
doc.add_paragraph('1. Monthly Rent: Rs. 25,000')
doc.add_paragraph('2. Security Deposit: Rs. 50,000')
doc.add_paragraph('3. Lease Duration: 12 months starting from January 1, 2024')

doc.add_heading('TENANT OBLIGATIONS:', level=1)
doc.add_paragraph('- Pay rent by 5th of every month')
doc.add_paragraph('- Maintain property in good condition')
doc.add_paragraph('- No subletting without written consent')

doc.add_heading('LANDLORD OBLIGATIONS:', level=1)
doc.add_paragraph('- Provide peaceful enjoyment of property')
doc.add_paragraph('- Maintain structural repairs')

doc.add_heading('TERMINATION:', level=1)
doc.add_paragraph('Either party may terminate with 30 days written notice.')

doc.add_paragraph('This agreement is governed by the laws of Maharashtra, India.')

doc.add_paragraph('Signatures:')
doc.add_paragraph('Landlord: _______________')
doc.add_paragraph('Tenant: _______________')
doc.add_paragraph('Date: _______________')

# Save the document
doc.save('test_rental_agreement.docx')
print("Created test_rental_agreement.docx successfully!")
